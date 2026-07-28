from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "article_revision" / "output"
TARGETS = [OUT / "artigo_revisado_final.docx", ROOT / "docs" / "Artigo Similaridade Facial.docx"]

def remove(p):
    p._element.getparent().remove(p._element)

for target in TARGETS:
    doc = Document(target)
    body = doc._element.body
    deleting = False
    for child in list(body):
        if child.tag.endswith("}p"):
            text = "".join(child.itertext())
            if text.startswith("3.6 Ensaio de sensibilidade"):
                deleting = True
            if deleting and text.startswith("4 Discussão"):
                deleting = False
        if deleting:
            body.remove(child)
    for p in list(doc.paragraphs):
        if p.text.startswith("Razão de enriquecimento") or p.text.startswith("Enriquecimento ="):
            remove(p)
        elif p.text.startswith("3.2 Concentração institucional"):
            p.text = "3.2 Lacuna de rastreabilidade da fonte"
        elif p.text.startswith("3.3 Comparação formal"):
            p.text = "3.3 Estados históricos e reprodução não controlável"
        elif p.text.startswith("3.4 Estabilidade em múltiplas sementes"):
            p.text = "3.4 Sensibilidades computacionais da partição"
        elif p.text.startswith("3.5 Ablação do escore"):
            p.text = "3.5 Circularidade e avaliação cross-fitted"
        elif p.text.startswith("4.3 Concentração por fonte"):
            p.text = "4.3 Lacuna de fonte e não identificabilidade"
    doc.save(target)

doc = Document(OUT / "artigo_revisado_final.docx")
lines = []
for p in doc.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    if p.style.name.startswith("Heading 1"):
        lines.append("# " + text)
    elif p.style.name.startswith("Heading 2"):
        lines.append("## " + text)
    else:
        lines.append(text)
    lines.append("")
(OUT / "artigo_revisado_final.md").write_text("\n".join(lines), encoding="utf-8")
