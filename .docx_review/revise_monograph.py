from __future__ import annotations

import re
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Henrique Lima\Desktop\Repositórios\ia-uspJailer")
SOURCE = Path(
    r"C:\Users\Henrique Lima\Downloads\SimilaridadeFacial_atualizado_analise_demografica_2026-08-16.docx"
)
OUTPUT = ROOT / "docs" / "Listas públicas de procurados não medem criminalidade — revisão metodológica.docx"
WORK = ROOT / ".docx_review"
FIGURES = WORK / "figures"
TABLES = ROOT / "research_audit_v2" / "outputs" / "final" / "tables"
DEMOGRAPHIC_TABLES = (
    ROOT / "research_audit_v2" / "outputs" / "demographic_composition" / "tables"
)

SKILL_SCRIPTS = Path(
    r"C:\Users\Henrique Lima\.codex\plugins\cache\openai-primary-runtime\documents\26.813.12317\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa  # noqa: E402


TITLE_PT = (
    "Listas Públicas De Procurados Não Medem Criminalidade: Auditoria sociotécnica da "
    "estabilidade e dos limites inferenciais de um pipeline de similaridade facial"
)
TITLE_EN = (
    "Public wanted-person lists do not measure criminality: a sociotechnical audit of "
    "stability and inferential limits in a face-similarity pipeline"
)

SUMMARY_PT = (
    "Esta monografia investiga uma falha recorrente em sistemas de inteligência artificial: "
    "o pipeline produz um rótulo a partir de uma representação, recupera esse rótulo com alta "
    "separabilidade e converte a circularidade em aparente evidência sobre pessoas. Propõe-se "
    "um framework de auditoria de construtos endógenos que distingue separabilidade interna, "
    "estabilidade computacional, validade de construto e validade externa. O caso empírico "
    "audita um pipeline de similaridade facial aplicado a listas públicas de procurados, que "
    "são produtos de seleção institucional e não amostras probabilísticas da criminalidade. "
    "Na reconstrução agrupada, 9.482 embeddings válidos foram avaliados por cross-fitting sem "
    "ajuste no teste. A ROC-AUC média foi 0,896, indicando recuperabilidade interna do alvo "
    "sintético, não mensuração de criminalidade. A análise de estabilidade reuniu 611 "
    "configurações e 60.045 comparações pareadas; em k = 64, as medianas foram 0,085 para ARI "
    "e 0,004 para Jaccard do alvo. Um controle sintético confirmou, em três demonstrações, que "
    "um alvo induzido pela geometria pode ser altamente separável e mudar quando o agrupamento "
    "é perturbado. Em braço independente com FairFace, os cenários demográficos apresentaram, "
    "contra a referência, ARI mediano de 0,182 a 0,199 e Jaccard próximo de 0,001, embora a "
    "prevalência do alvo e a ROC-AUC variassem pouco. Como a identidade das imagens e a "
    "inicialização também variam, o experimento não identifica efeito causal demográfico. Os "
    "resultados demonstram a falácia da separabilidade endógena: recuperar um rótulo produzido "
    "pelo próprio sistema não estabelece a existência nem a validade de um construto externo."
)

ABSTRACT_EN = (
    "This monograph examines a recurring failure in artificial-intelligence systems: a "
    "pipeline produces a label from a representation, recovers that label with high "
    "separability, and converts this circularity into apparent evidence about people. It "
    "proposes an audit framework for endogenous constructs that distinguishes internal "
    "separability, computational stability, construct validity, and external validity. The "
    "empirical case audits a face-similarity pipeline applied to public wanted-person lists, "
    "which are products of institutional selection rather than probability samples of "
    "criminality. In the grouped reconstruction, 9,482 valid embeddings were evaluated by "
    "cross-fitting with no test-set fitting. Mean ROC-AUC was 0.896, indicating internal "
    "recoverability of the synthetic target rather than measurement of criminality. The "
    "stability analysis comprised 611 configurations and 60,045 pairwise comparisons; at "
    "k = 64, median ARI was 0.085 and median target Jaccard was 0.004. A synthetic control "
    "confirmed in three demonstrations that a geometry-induced target can be highly separable "
    "and change when clustering is perturbed. In an independent FairFace arm, demographic "
    "scenarios yielded median ARI values from 0.182 to 0.199 and target Jaccard near 0.001 "
    "against the reference, while target prevalence and ROC-AUC changed little. Because image "
    "membership and initialization also vary, the experiment does not identify a causal "
    "demographic effect. The results demonstrate the endogenous-separability fallacy: "
    "recovering a label produced by the system itself does not establish the existence or "
    "validity of an external construct."
)


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\wÀ-ÿ'-]+\b", text))


def find_paragraph(doc: Document, prefix: str, style: str | None = None) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix) and (
            style is None or paragraph.style.name == style
        ):
            return paragraph
    raise KeyError(f"Paragraph not found: {prefix!r}, style={style!r}")


def set_text(paragraph: Paragraph, text: str) -> Paragraph:
    paragraph.clear()
    paragraph.add_run(text)
    return paragraph


def set_labeled_text(paragraph: Paragraph, label: str, text: str) -> Paragraph:
    paragraph.clear()
    lead = paragraph.add_run(label)
    lead.bold = True
    paragraph.add_run(text)
    return paragraph


def new_paragraph_after(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def new_paragraph_before(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def delete_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def delete_table(table) -> None:
    parent = table._tbl.getparent()
    if parent is not None:
        parent.remove(table._tbl)


def next_table_after(paragraph: Paragraph):
    node = paragraph._p.getnext()
    while node is not None:
        if node.tag == qn("w:tbl"):
            from docx.table import Table

            return Table(node, paragraph._parent)
        if node.tag == qn("w:p"):
            candidate = Paragraph(node, paragraph._parent)
            if candidate.text.strip():
                break
        node = node.getnext()
    raise KeyError(f"No table found after paragraph {paragraph.text!r}")


def next_text_paragraph(paragraph: Paragraph) -> Paragraph:
    node = paragraph._p.getnext()
    while node is not None:
        if node.tag == qn("w:p"):
            candidate = Paragraph(node, paragraph._parent)
            if candidate.text.strip():
                return candidate
        node = node.getnext()
    raise KeyError(f"No following text paragraph after {paragraph.text!r}")


def insert_picture_after(anchor: Paragraph, path: Path, width_inches: float = 6.1) -> Paragraph:
    paragraph = new_paragraph_after(anchor, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(path), width=Inches(width_inches))
    return paragraph


def set_no_hyphenation(paragraph: Paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    node = ppr.find(qn("w:suppressAutoHyphens"))
    if node is None:
        node = OxmlElement("w:suppressAutoHyphens")
        ppr.append(node)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        if edge in {"top", "bottom"}:
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "8")
            node.set(qn("w:color"), "5B6B73")
        elif edge == "insideH":
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "3")
            node.set(qn("w:color"), "C7D0D5")
        else:
            node.set(qn("w:val"), "nil")


def format_table(table, weights: list[float], font_size: float = 9.0) -> None:
    total = section_content_width_dxa(table._parent.part.document.sections[-1]) - 120
    widths = column_widths_from_weights(weights, total)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=total,
        indent_dxa=120,
        cell_margins_dxa={"top": 95, "bottom": 95, "start": 120, "end": 120},
    )
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            trpr = row._tr.get_or_add_trPr()
            repeat = trpr.find(qn("w:tblHeader"))
            if repeat is None:
                repeat = OxmlElement("w:tblHeader")
                trpr.append(repeat)
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                set_cell_shading(cell, "E8EEF1")
            for paragraph in cell.paragraphs:
                set_no_hyphenation(paragraph)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.right_indent = Pt(0)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                if row_idx == 0 or len(paragraph.text) < 22:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    if row_idx == 0:
                        run.bold = True


def replace_table_after_caption(
    doc: Document,
    caption: Paragraph,
    headers: list[str],
    rows: list[list[str]],
    weights: list[float],
    font_size: float = 8.5,
):
    old_table = next_table_after(caption)
    delete_table(old_table)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = value
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    caption._p.addnext(table._tbl)
    format_table(table, weights, font_size)
    return table


def make_pipeline_figure(path: Path) -> None:
    image = Image.new("RGB", (1900, 830), "white")
    draw = ImageDraw.Draw(image)
    regular = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 30)
    small = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 27)
    bold = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 29)
    stages = [
        (70, 160, "Seleção\ninstitucional", "#F4D7B2"),
        (535, 160, "Publicação\ne vigência", "#F4D7B2"),
        (1000, 160, "Coleta", "#F4D7B2"),
        (1465, 160, "Detecção e\nalinhamento", "#F4D7B2"),
        (1465, 500, "Embedding\nZ = f(X)", "#CFE2F3"),
        (1000, 500, "Clustering\nY = h(Z)", "#CFE2F3"),
        (535, 500, "Escolha do\nalvo", "#CFE2F3"),
        (70, 500, "Score e\navaliação g(Z)", "#CFE2F3"),
    ]
    boxes: list[tuple[int, int, int, int]] = []
    for x, y, label, color in stages:
        box = (x, y, x + 350, y + 135)
        draw.rounded_rectangle(box, radius=18, fill=color, outline="#40515A", width=3)
        draw.multiline_text(
            ((box[0] + box[2]) / 2, (box[1] + box[3]) / 2),
            label,
            font=regular,
            fill="#25343B",
            anchor="mm",
            align="center",
            spacing=3,
        )
        boxes.append(box)

    def arrow(start: tuple[float, float], end: tuple[float, float]) -> None:
        draw.line((start, end), fill="#40515A", width=5)
        x2, y2 = end
        x1, y1 = start
        angle = np.arctan2(y2 - y1, x2 - x1)
        length = 18
        wing = 0.55
        p1 = (x2 - length * np.cos(angle - wing), y2 - length * np.sin(angle - wing))
        p2 = (x2 - length * np.cos(angle + wing), y2 - length * np.sin(angle + wing))
        draw.polygon([end, p1, p2], fill="#40515A")

    arrow_pairs = [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5), (5, 6), (6, 7)]
    for a, b in arrow_pairs:
        xa1, ya1, xa2, ya2 = boxes[a]
        xb1, yb1, xb2, yb2 = boxes[b]
        if a == 3:
            arrow(((xa1 + xa2) / 2, ya2), ((xb1 + xb2) / 2, yb1))
        elif ya1 == yb1 and xb1 > xa1:
            arrow((xa2, (ya1 + ya2) / 2), (xb1, (yb1 + yb2) / 2))
        else:
            arrow((xa1, (ya1 + ya2) / 2), (xb2, (yb1 + yb2) / 2))
    draw.text(
        (950, 70),
        "Seleção e observabilidade: condicionam o universo S = 1",
        anchor="mm",
        font=bold,
        fill="#8A4F08",
    )
    draw.text(
        (950, 705),
        "Geometria computacional: separabilidade e estabilidade internas",
        anchor="mm",
        font=bold,
        fill="#1D5A82",
    )
    draw.text(
        (950, 785),
        "Sem vínculo observável entre Y e um construto externo C, recuperabilidade não estabelece validade de construto.",
        anchor="mm",
        font=small,
        fill="#7A1E1E",
    )
    image.save(path, quality=96)


def make_paradox_figure(path: Path) -> None:
    summary = pd.read_csv(TABLES / "stability_summary.csv")
    folds = pd.read_csv(TABLES / "cross_fitted_metrics.csv")
    stoch = summary[summary["instability_type"].eq("stochastic")]
    ks = sorted(stoch["k"].unique())
    image = Image.new("RGB", (1900, 1320), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 25)
    small = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 21)
    bold = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 28)
    title_font = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 38)
    draw.text((950, 45), "Paradoxo separabilidade–estabilidade", font=title_font, fill="#24353E", anchor="mm")
    panels = [(85, 115, 900, 625), (1000, 115, 1815, 625), (85, 700, 900, 1210), (1000, 700, 1815, 1210)]

    def axes_box(box, heading, y_max):
        x1, y1, x2, y2 = box
        draw.rounded_rectangle(box, radius=16, outline="#C7D0D5", width=2, fill="#FCFDFD")
        draw.text(((x1 + x2) / 2, y1 + 34), heading, font=bold, fill="#24353E", anchor="mm")
        plot = (x1 + 80, y1 + 78, x2 - 35, y2 - 65)
        px1, py1, px2, py2 = plot
        draw.line((px1, py1, px1, py2), fill="#59676E", width=2)
        draw.line((px1, py2, px2, py2), fill="#59676E", width=2)
        for frac in np.linspace(0, 1, 5):
            yy = py2 - frac * (py2 - py1)
            draw.line((px1, yy, px2, yy), fill="#E6EAEC", width=1)
            draw.text((px1 - 10, yy), f"{frac*y_max:.2f}", font=small, fill="#59676E", anchor="rm")
        return plot

    plot = axes_box(panels[0], "A. Recuperabilidade interna por dobra", 1.0)
    px1, py1, px2, py2 = plot
    labels = ["ROC-AUC", "PR-AUC", "Baseline PR"]
    cols = ["roc_auc", "pr_auc", "pr_auc_baseline"]
    colors = ["#2E6F9E", "#4E9A8D", "#C8A64B"]
    centers = np.linspace(px1 + 120, px2 - 120, 3)
    for idx, (label, column, color) in enumerate(zip(labels, cols, colors)):
        value = float(folds[column].mean())
        cx = centers[idx]
        top = py2 - value * (py2 - py1)
        draw.rectangle((cx - 55, top, cx + 55, py2), fill=color)
        for dot_x, dot in zip(np.linspace(cx - 35, cx + 35, len(folds)), folds[column]):
            yy = py2 - float(dot) * (py2 - py1)
            draw.ellipse((dot_x - 5, yy - 5, dot_x + 5, yy + 5), fill="#24353E")
        draw.text((cx, py2 + 26), label, font=small, fill="#24353E", anchor="mm")
        draw.text((cx, top - 16), f"{value:.3f}", font=small, fill="#24353E", anchor="mm")

    def metric_panel(box, metric, heading, color, y_max):
        plot = axes_box(box, heading, y_max)
        px1, py1, px2, py2 = plot
        rows = stoch[stoch["metric"].eq(metric)].set_index("k").loc[ks]
        xs = np.linspace(px1 + 25, px2 - 25, len(ks))

        def yy(value):
            return py2 - (float(value) / y_max) * (py2 - py1)

        p05 = [(x, yy(v)) for x, v in zip(xs, rows["p05"])]
        p95 = [(x, yy(v)) for x, v in zip(xs, rows["p95"])]
        q1 = [(x, yy(v)) for x, v in zip(xs, rows["q1"])]
        q3 = [(x, yy(v)) for x, v in zip(xs, rows["q3"])]
        median = [(x, yy(v)) for x, v in zip(xs, rows["median"])]
        draw.polygon(p05 + list(reversed(p95)), fill="#E8E4ED")
        draw.polygon(q1 + list(reversed(q3)), fill="#D3C7DE")
        draw.line(median, fill=color, width=5, joint="curve")
        for (x, y), k in zip(median, ks):
            draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill=color)
            draw.text((x, py2 + 25), str(k), font=small, fill="#24353E", anchor="mm")
        draw.text(((px1 + px2) / 2, py2 + 52), "Número de clusters (k)", font=small, fill="#59676E", anchor="mm")
        draw.text((px2 - 150, py1 + 20), "P05–P95   IQR   Mediana", font=small, fill="#59676E", anchor="mm")

    metric_panel(panels[1], "ari", "B. Concordância das partições (ARI)", "#B45252", 0.16)
    metric_panel(panels[2], "target_jaccard", "C. Sobreposição do conjunto-alvo", "#7B5AA6", 0.46)
    metric_panel(panels[3], "target_prevalence", "D. Prevalência do maior cluster", "#3E8C65", 0.09)
    draw.text(
        (950, 1285),
        "Pontos no painel A representam as cinco dobras; painéis B–D resumem 100 sementes por k.",
        font=small,
        fill="#4C5A61",
        anchor="mm",
    )
    image.save(path, quality=96)


def add_reference_before(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = new_paragraph_before(anchor, text, "Reference")
    return paragraph


def build() -> None:
    WORK.mkdir(parents=True, exist_ok=True)
    FIGURES.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    pipeline_figure = FIGURES / "pipeline_sociotecnico.png"
    paradox_figure = FIGURES / "paradoxo_separabilidade_estabilidade.png"
    make_pipeline_figure(pipeline_figure)
    make_paradox_figure(paradox_figure)

    doc = Document(SOURCE)
    original_tables = list(doc.tables)

    # Refresh the cataloguing card without tying it to a page count that changes
    # whenever the analytical revision is repaginated.
    catalog_replacements = {
        "A falsa ideia de “criminalidade facial”: viés racial, listas públicas de procurados e os limites da similaridade facial: Auditoria sociotécnica de um pipeline que produz separabilidade sem medir criminalidade": TITLE_PT,
        "1 recurso online (31 p.) : il., quadros, tabelas.": "1 recurso online: il., quadros, tabelas.",
        "Inclui referências e material suplementar.": "Inclui referências e apêndice.",
    }
    for cell in original_tables[0]._cells:
        for paragraph in cell.paragraphs:
            updated = paragraph.text
            for old, new in catalog_replacements.items():
                updated = updated.replace(old, new)
            if updated != paragraph.text:
                set_text(paragraph, updated)

    # Remove the article-style callouts that duplicate the abstract/results.
    delete_table(original_tables[1])
    delete_table(original_tables[9])

    # Title, title-page metadata, and bibliographic citation.
    set_text(find_paragraph(doc, "A falsa ideia", "FM Subtítulo"), "LISTAS PÚBLICAS DE PROCURADOS NÃO MEDEM CRIMINALIDADE")
    subtitle_cover = [
        p for p in doc.paragraphs if p.style.name == "FM Subtítulo" and p.text.startswith("Viés racial")
    ][0]
    set_text(
        subtitle_cover,
        "Auditoria sociotécnica da estabilidade e dos limites inferenciais de um pipeline de similaridade facial",
    )
    english_cover = [p for p in doc.paragraphs if p.style.name == "FM Subtítulo" and p.text.startswith("The false")][0]
    set_text(english_cover, TITLE_EN)
    title_meta = find_paragraph(doc, "Título:")
    set_text(title_meta, f"Título: {TITLE_PT}")
    cite = find_paragraph(doc, "GUSMÃO, Henrique Lima.", "TechnicalText")
    set_text(
        cite,
        f"GUSMÃO, Henrique Lima. {TITLE_PT}. 1. ed. Belo Horizonte: edição do autor, 2026. "
        "Publicação eletrônica depositada no Zenodo. DOI: 10.5281/zenodo.21659240. "
        "ISBN: 978-65-02-26124-8.",
    )

    # Move resumo and abstract before the table of contents; remove the duplicate article header.
    toc_heading = find_paragraph(doc, "SUMÁRIO", "FrontHeading")
    summary_heading = find_paragraph(doc, "RESUMO", "Heading 1")
    summary_text = next_text_paragraph(summary_heading)
    summary_keywords = find_paragraph(doc, "Palavras-chave:")
    abstract_heading = find_paragraph(doc, "ABSTRACT", "Heading 1")
    abstract_text = next_text_paragraph(abstract_heading)
    abstract_keywords = find_paragraph(doc, "Keywords:")
    set_text(summary_text, SUMMARY_PT)
    summary_heading.style = "FrontHeading"
    set_text(
        summary_keywords,
        "Palavras-chave: validade de construto; viés de seleção; reconhecimento facial; "
        "auditoria algorítmica; fisiognomia algorítmica.",
    )
    set_text(abstract_text, ABSTRACT_EN)
    abstract_heading.style = "FrontHeading"
    set_text(
        abstract_keywords,
        "Keywords: construct validity; selection bias; face recognition; algorithmic auditing; "
        "algorithmic physiognomy.",
    )
    for paragraph in (
        summary_heading,
        summary_text,
        summary_keywords,
        abstract_heading,
        abstract_text,
        abstract_keywords,
    ):
        toc_heading._p.addprevious(paragraph._p)
    summary_heading.paragraph_format.page_break_before = True
    abstract_heading.paragraph_format.page_break_before = True
    toc_heading.paragraph_format.page_break_before = False
    toc_page_break = new_paragraph_before(toc_heading, "", "Normal")
    toc_page_break.add_run().add_break(WD_BREAK.PAGE)
    for paragraph in list(doc.paragraphs):
        if paragraph.style.name in {"TOC 1 Custom", "TOC 2 Custom", "TOC 3 Custom"}:
            delete_paragraph(paragraph)
        elif paragraph.style.name in {"Title", "Subtitle", "Author", "Affiliation", "PublicationMeta"}:
            delete_paragraph(paragraph)
    toc_marker = new_paragraph_after(toc_heading, "[[TOC]]", "Normal")
    toc_marker.paragraph_format.space_after = Pt(0)

    # Rewrite the introduction around the broader methodological contribution.
    intro = find_paragraph(doc, "1 INTRODUÇÃO", "Heading 1")
    # A preserved section break already starts the textual body on a new page.
    intro.paragraph_format.page_break_before = False
    p91 = find_paragraph(doc, "O ponto de partida desta monografia")
    set_text(
        p91,
        "Esta monografia investiga como uma regularidade produzida por um sistema de inteligência "
        "artificial pode adquirir aparência de descoberta sobre seres humanos. O caso de listas "
        "públicas de procurados é deliberadamente crítico: fotografias institucionalmente "
        "selecionadas são transformadas em embeddings, agrupadas e pontuadas, embora o corpus não "
        "contenha uma variável observável que represente criminalidade. A questão central deixa de "
        "ser se um rosto revela crime e passa a ser como detectar a conversão de circularidade "
        "computacional em alegação substantiva.",
    )
    p92 = find_paragraph(doc, "Listas públicas de procurados")
    set_text(
        p92,
        "Listas públicas de procurados registram pessoas que determinadas instituições decidiram "
        "localizar e divulgar em certo momento. Não são levantamentos probabilísticos de delitos, "
        "autores ou populações. Na Interpol, a Red Notice é um pedido de localização e prisão "
        "provisória; apenas parte dos avisos se torna pública, e a observação depende de jurisdição, "
        "política de divulgação, vigência, disponibilidade de fotografia e capacidade de coleta "
        "(Interpol, 2026a, 2026b). Esse encadeamento seleciona o universo observado antes de qualquer "
        "modelo.",
    )
    p93 = find_paragraph(doc, "A questão racial")
    set_text(
        p93,
        "A dimensão racial torna o problema especialmente sensível. Desigualdades documentadas no "
        "sistema penal e em procedimentos de reconhecimento coexistem com bases biométricas cuja "
        "composição e qualidade afetam os erros. No Brasil, 55,5% da população se autodeclarou preta "
        "ou parda no Censo 2022; estatísticas penais e recomendações do CNJ também registram "
        "disparidades relevantes (IBGE, 2022; Conselho Nacional de Justiça, 2022; Secretaria "
        "Nacional de Políticas Penais, 2025). Esses dados contextualizam o risco sociotécnico, mas "
        "não autorizam atribuir raça/cor às imagens do corpus nem estimar seletividade racial sem "
        "metadados por registro e denominadores compatíveis.",
    )
    p94 = find_paragraph(doc, "Em aprendizagem de máquina")
    set_text(
        p94,
        "O problema formal aparece quando o pipeline deriva um alvo Y de uma representação Z = f(X) "
        "e avalia um score g(Z) contra esse mesmo alvo. Uma AUC alta pode então demonstrar apenas "
        "recuperabilidade da regra construída. Chama-se aqui falácia da separabilidade endógena a "
        "passagem indevida da recuperabilidade de Y para a existência ou validade de um construto "
        "externo C. A auditoria deve separar quatro perguntas: o alvo é separável, é estável, mede o "
        "construto alegado e se generaliza para uma população definida?",
    )
    p95 = find_paragraph(doc, "Esta monografia audita")
    set_text(
        p95,
        "A monografia audita o estado preservado e a reconstrução final do repositório "
        "are-you-a-criminal-ML. O desenho combina análise de identificabilidade, linhagem de dados, "
        "cross-fitting agrupado, controles negativos, estabilidade sob perturbações e um braço "
        "externo com FairFace. O objetivo não é construir um detector de criminalidade, mas oferecer "
        "um protocolo reutilizável para reconhecer quando um sistema transforma seus próprios "
        "rótulos em aparente evidência científica.",
    )
    rq = find_paragraph(doc, "As perguntas de pesquisa")
    set_text(
        rq,
        "As perguntas de pesquisa são: RQ1 — quais decisões institucionais e técnicas condicionam "
        "o corpus e seus limites externos? RQ2 — qual é a separabilidade interna do alvo endógeno? "
        "RQ3 — quão estáveis são a partição e a identidade do alvo sob mudanças de seed, k, ordem, "
        "minibatch e representação? RQ4 — quais alegações sobre criminalidade, raça e populações "
        "permanecem não identificáveis?",
    )
    contribution_heading = new_paragraph_after(rq, "1.1 Contribuições", "Heading 2")
    cursor = contribution_heading
    contributions = [
        ("C1 — contribuição conceitual. ", "Distingue separabilidade, estabilidade, validade de construto e validade externa."),
        ("C2 — contribuição metodológica. ", "Formaliza uma matriz de identificabilidade e verificabilidade para alegações sociotécnicas."),
        ("C3 — contribuição computacional. ", "Integra linhagem, auditoria de vazamento, cross-fitting agrupado, controles negativos e estabilidade pairwise."),
        ("C4 — contribuição empírica. ", "Mostra que métricas agregadas podem permanecer altas enquanto a partição e a identidade do alvo variam substancialmente."),
        ("C5 — contribuição sociotécnica. ", "Trata seleção institucional, publicação, coleta e processamento como partes do sistema avaliado."),
    ]
    for label, text in contributions:
        cursor = new_paragraph_after(cursor, style="Normal")
        set_labeled_text(cursor, label, text)
        cursor.paragraph_format.left_indent = Inches(0.2)
        cursor.paragraph_format.first_line_indent = Inches(-0.2)
    related_heading = new_paragraph_after(cursor, "1.2 Trabalhos relacionados e fundamentação", "Heading 2")
    cursor = new_paragraph_after(
        related_heading,
        "O enquadramento articula seis literaturas: crítica à fisiognomia algorítmica; validade de "
        "construto e mensuração; viés de seleção; auditoria de datasets e modelos; reconhecimento "
        "facial e diferenças demográficas; e associações entre aparência e categorias sociais ou "
        "morais em modelos multimodais. Bowyer et al. (2020) demonstram a fragilidade da suposta "
        "‘criminalidade pela face’; Jacobs e Wallach (2021) mostram que mensuração e fairness "
        "dependem da validade do construto; Raji et al. (2020, 2022) situam auditoria e alegações de "
        "funcionalidade ao longo do ciclo de vida do sistema.",
        "Normal",
    )
    cursor = new_paragraph_after(
        cursor,
        "O problema permanece atual em modelos multimodais. Birhane et al. (2024) avaliaram 14 "
        "modelos visiolinguísticos e observaram associações racializadas com a classe ‘criminal’, "
        "cujo comportamento mudou com a escala do modelo e do conjunto de treinamento. O resultado "
        "não equivale ao experimento desta monografia, mas reforça que rótulos morais podem emergir "
        "de escolhas de dados, classes e representação, exigindo auditoria de sua proveniência e "
        "significado.",
        "Normal",
    )
    new_paragraph_after(
        cursor,
        "Na biometria, a avaliação contínua do NIST destaca que diferenças demográficas de erro "
        "dependem do algoritmo, da tarefa e da qualidade da imagem (NIST, 2026). FairFace foi criado "
        "para estudar desequilíbrios de atributos em um domínio próprio e fornece categorias "
        "históricas de anotação, não uma ontologia biológica nem autoidentificação irrestrita "
        "(Kärkkäinen; Joo, 2021). Por isso, o braço externo é tratado como análise de sensibilidade de "
        "composição, sem transferência de rótulos ao corpus principal.",
        "Normal",
    )

    # Insert the formal framework and renumber methods subsections.
    original_22 = find_paragraph(doc, "2.2 Classes de identificabilidade", "Heading 2")
    framework_heading = new_paragraph_before(
        original_22, "2.2 Framework de auditoria da separabilidade endógena", "Heading 2"
    )
    cursor = new_paragraph_after(
        framework_heading,
        "Seja X o dado observado, Z = f(X) uma representação, Y = h(Z) um alvo produzido ou "
        "selecionado pelo pipeline e s = g(Z) um score avaliado contra Y. A falácia ocorre quando o "
        "sucesso de s em recuperar Y é apresentado como evidência de que Y mede um construto externo "
        "C. A implicação recuperabilidade(Y) ⇒ validade(C) não é lógica nem empiricamente garantida.",
        "Normal",
    )
    equation = new_paragraph_after(
        cursor,
        "Y = h(f(X));  s = g(f(X));  separabilidade(s, Y) ⇏ existência ou validade de C",
        "TechnicalText",
    )
    equation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(6)
    equation.paragraph_format.space_after = Pt(6)
    cursor = new_paragraph_after(
        equation,
        "O framework audita quatro planos. (1) Separabilidade interna: ROC-AUC e PR-AUC respondem "
        "quão bem o score ordena o alvo endógeno. (2) Estabilidade: ARI, Jaccard e prevalência "
        "respondem quanto a partição e o alvo dependem de escolhas computacionais. (3) Validade de "
        "construto: requer evidência de que Y representa C, e não apenas a geometria que o gerou. "
        "(4) Validade externa: requer população-alvo, denominadores, seleção observável e condições de "
        "generalização. Cada plano bloqueia uma inferência distinta.",
        "Normal",
    )
    new_paragraph_after(
        cursor,
        "O protocolo operacional consiste em localizar a origem de Y; separar ajuste, seleção e "
        "avaliação; executar controles que rompam a relação geométrica; perturbar sementes, "
        "hiperparâmetros, ordem e representação; e classificar como não identificável toda alegação "
        "externa sem variável, comparador ou denominador observável. O caso de ‘criminalidade "
        "facial’ funciona como demonstração construtiva dessa falha, mas o protocolo também se aplica "
        "a inferências de personalidade, empregabilidade, periculosidade, emoção ou orientação "
        "política.",
        "Normal",
    )
    method_renames = {
        "2.2 Classes de identificabilidade": "2.3 Classes de identificabilidade",
        "2.3 Cadeia sociotécnica de seleção e status da evidência": "2.4 Cadeia sociotécnica de seleção e status da evidência",
        "2.4 Proveniência de fonte e análises excluídas": "2.5 Proveniência de fonte e análises excluídas",
        "2.5 Análise de estabilidade computacional": "2.6 Análise de estabilidade computacional",
        "2.6 Avaliação agrupada e contrato de reprodução": "2.7 Avaliação agrupada e contrato de reprodução",
        "2.6.1 Construção e limites do group_id": "2.7.1 Construção e limites do group_id",
        "2.7 Níveis de verificabilidade computacional": "2.8 Níveis de verificabilidade computacional",
        "2.8 Experimento externo de composição demográfica": "2.9 Experimento externo de composição demográfica",
    }
    for old, new in method_renames.items():
        set_text(find_paragraph(doc, old), new)

    chain = find_paragraph(doc, "A inclusão final foi representada")
    pipeline_caption = new_paragraph_after(
        chain,
        "Figura 1 — Cadeia sociotécnica, endogeneidade do alvo e limites inferenciais.",
        "TableCaption",
    )
    pipeline_image = insert_picture_after(pipeline_caption, pipeline_figure, 6.1)
    new_paragraph_after(
        pipeline_image,
        "Fonte: elaboração própria. O diagrama distingue seleção/observabilidade, geometria interna e "
        "o vínculo externo que precisaria ser demonstrado.",
        "TableSource",
    )

    # Clarify the external experiment and the exact reproducible implementation boundary.
    fairface_intro = find_paragraph(doc, "Como o corpus principal não possui")
    set_text(
        fairface_intro,
        "Como o corpus principal não possui raça/cor documentada nem vínculo verificável com bases "
        "demográficas externas, a sensibilidade à composição foi examinada em braço independente com "
        "FairFace (Kärkkäinen; Joo, 2021). Foram usadas apenas as sete categorias fornecidas pela base; "
        "nenhuma categoria foi inferida visualmente ou transferida ao pipeline de procurados. Cada "
        "cenário contém 36.456 registros sem reposição. As proporções foram controladas, mas a "
        "identidade específica das imagens também varia entre cenários; portanto, o desenho não isola "
        "o efeito da proporção demográfica do efeito da amostragem de registros.",
    )
    fairface_impl = find_paragraph(doc, "A liberação `margin025`")
    set_text(
        fairface_impl,
        "A liberação margin025, recortada e alinhada upstream com dlib.get_face_chip(), foi fornecida "
        "diretamente ao reconhecedor do pacote InsightFace buffalo_l. A implementação usa InsightFace "
        "1.0.1, ONNX Runtime 1.18.1, leitura RGB com conversão para BGR, normalização de pixels e "
        "redimensionamento executados pelo reconhecedor, embedding de 512 dimensões e normalização L2 "
        "posterior. O arquivo de reconhecimento disponível no ambiente é w600k_r50.onnx, mas seu hash "
        "não foi incorporado ao manifesto da corrida; por isso, o SHA-256 observado posteriormente não "
        "é atribuído retroativamente à execução. MiniBatchKMeans, n_init = 3, batch_size = 1.024, "
        "max_iter = 100, regra do maior cluster, score cosseno e avaliação foram mantidos constantes.",
    )
    fairface_thresholds = find_paragraph(doc, "As comparações entre cenários usam")
    set_text(
        fairface_thresholds,
        "As comparações entre cenários usam apenas a interseção dos identificadores presentes em "
        "ambos. Os limites ARI < 0,90, Jaccard < 0,80, |Δ prevalência| ≥ 0,02 e |Δ AUC| ≥ 0,03 foram "
        "pré-declarados como heurísticas descritivas para organizar a leitura, não como limiares de "
        "significância ou relevância substantiva validados externamente. A inferência principal usa "
        "distribuições e tamanhos de efeito contínuos.",
    )

    # Replace the stability scope table with the actual distributional result.
    stability_caption = find_paragraph(doc, "Tabela 2 —")
    set_text(
        stability_caption,
        "Tabela 2 — Estabilidade estocástica por k: medianas, IQR e intervalos P05–P95.",
    )
    stability = pd.read_csv(TABLES / "stability_summary.csv")
    stoch = stability[stability["instability_type"].eq("stochastic")]
    table_rows: list[list[str]] = []
    for k in sorted(stoch.k.unique()):
        frame = stoch[stoch.k.eq(k)].set_index("metric")
        ari = frame.loc["ari"]
        jac = frame.loc["target_jaccard"]
        prev = frame.loc["target_prevalence"]
        table_rows.append(
            [
                str(int(k)),
                f"{ari['median']:.3f} ({ari['q1']:.3f}–{ari['q3']:.3f})",
                f"{ari['p05']:.3f}–{ari['p95']:.3f}",
                f"{jac['median']:.3f} ({jac['q1']:.3f}–{jac['q3']:.3f})",
                f"{jac['p05']:.3f}–{jac['p95']:.3f}",
                f"{prev['median']:.3f} ({prev['q1']:.3f}–{prev['q3']:.3f})",
            ]
        )
    replace_table_after_caption(
        doc,
        stability_caption,
        ["k", "ARI mediana (IQR)", "ARI P05–P95", "Jaccard mediana (IQR)", "Jaccard P05–P95", "Prevalência mediana (IQR)"],
        table_rows,
        [0.45, 1.45, 1.05, 1.55, 1.05, 1.55],
        8.0,
    )
    stability_source = find_paragraph(doc, "Fonte: manifesto e relatórios")
    set_text(
        stability_source,
        "Fonte: stability_summary.csv. Para ARI e Jaccard, n = 4.950 pares distintos por k; para "
        "prevalência, n = 100 execuções por k. IQR = Q1–Q3.",
    )
    st_p1 = find_paragraph(doc, "A edição anterior resumia oito contrastes")
    set_text(
        st_p1,
        "A corrida final torna visível o resultado das 60.045 comparações. Entre sementes, a mediana "
        "do ARI diminuiu de 0,113 em k = 32 para 0,070 em k = 128. A mediana do Jaccard do conjunto-alvo "
        "permaneceu entre 0,001 e 0,051; em k = 64 foi 0,004 (IQR 0,000–0,047; P05–P95 0,000–0,126). "
        "Esses valores mostram baixa concordância tanto da partição completa quanto da identidade do "
        "maior cluster na região examinada.",
    )
    st_p2 = find_paragraph(doc, "Assim, a conclusão cientificamente sustentada")
    set_text(
        st_p2,
        "A sensibilidade não se restringiu à semente. Em k = 64, mudanças de ordem produziram ARI "
        "mediano 0,100 e Jaccard 0,000; mudanças de batch_size, ARI 0,085 e Jaccard 0,003; o contraste "
        "original L2 × PCA-64 resultou em ARI 0,090 e Jaccard 0,000. A prevalência do alvo também "
        "variou: entre batch_size, a mediana foi 0,040 e o intervalo observado 0,028–0,125. Como o "
        "desenho cobre uma grade finita, os resultados sustentam instabilidade condicionada às "
        "perturbações executadas, não uma afirmação universal sobre todo algoritmo de clustering.",
    )

    # Report fold uncertainty without inventing an observation-level confidence interval.
    metric_result = find_paragraph(doc, "Na reconstrução atual, a ROC-AUC")
    set_text(
        metric_result,
        "Na reconstrução atual, a ROC-AUC média foi 0,896, com variação entre dobras de 0,882 a 0,904. "
        "A PR-AUC média foi 0,348 (0,211–0,448), frente a baseline média 0,054, e o Brier médio foi "
        "0,043 (0,034–0,049). Os arquivos finais preservam métricas agregadas por cinco dobras, mas "
        "não predições por observação; por isso, não se apresenta um intervalo bootstrap que não possa "
        "ser reconstruído de modo verificável. Esses números quantificam recuperabilidade interna do "
        "alvo sintético, não desempenho para criminalidade, culpa, raça/cor ou outro atributo externo.",
    )
    perf_table = next_table_after(find_paragraph(doc, "Tabela 3 —"))
    replacements = {
        "0,896170": "0,896",
        "0,348207": "0,348",
        "0,053890": "0,054",
        "0,043299": "0,043",
    }
    for row in perf_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                updated = paragraph.text
                for old, new in replacements.items():
                    updated = updated.replace(old, new)
                if updated != paragraph.text:
                    set_text(paragraph, updated)

    # Turn the redundant synthesis section into the synthetic-control result and add the central figure.
    synthesis_heading = find_paragraph(doc, "3.5 Síntese do aparente paradoxo", "Heading 2")
    set_text(synthesis_heading, "3.5 Controle sintético e paradoxo separabilidade–estabilidade")
    synthetic_text = new_paragraph_after(
        synthesis_heading,
        "O controle sintético passou nas três demonstrações pré-definidas. O clustering gerou um alvo "
        "sem variável social externa; um score derivado da mesma geometria recuperou esse alvo com "
        "ROC-AUC = 1,000; e a perturbação de k = 3 para k = 4 mudou completamente o conjunto-alvo "
        "(Jaccard = 0,000). O controle não representa pessoas nem prova um efeito no corpus principal. "
        "Ele isola o mecanismo lógico: separabilidade perfeita pode coexistir com ausência de validade "
        "externa e com identidade instável do alvo.",
        "Normal",
    )
    paradox_caption = new_paragraph_after(
        synthetic_text,
        "Figura 2 — Paradoxo separabilidade–estabilidade na reconstrução final.",
        "TableCaption",
    )
    paradox_image = insert_picture_after(paradox_caption, paradox_figure, 6.1)
    new_paragraph_after(
        paradox_image,
        "Fonte: elaboração própria a partir de cross_fitted_metrics.csv e stability_summary.csv. "
        "As métricas do painel A avaliam o alvo endógeno; os demais painéis resumem perturbações de seed e k.",
        "TableSource",
    )

    # Renumber and tighten the FairFace result language.
    set_text(find_paragraph(doc, "Figura 1 — Composição média"), "Figura 3 — Composição média do cluster-alvo por categoria fornecida pelo FairFace, em k = 64 e ao longo das sementes.")
    set_text(find_paragraph(doc, "Figura 2 — Estabilidade entre"), "Figura 4 — Estabilidade entre sementes no experimento externo, em k = 64.")
    ff_result = find_paragraph(doc, "O braço externo acrescenta")
    set_text(
        ff_result,
        "O braço externo mostra que cenários com proporções demográficas diferentes podem produzir "
        "estruturas materialmente distintas. Em k = 64, B, C e D apresentaram, contra A, ARI mediano "
        "de 0,182, 0,199 e 0,189 e Jaccard mediano do alvo próximo de 0,001. Esses contrastes incluem, "
        "contudo, tanto a mudança de proporção quanto a mudança dos registros amostrados e ocorrem "
        "sobre forte instabilidade entre sementes; não identificam efeito causal demográfico.",
    )

    # Refocus discussion on the reusable framework and reduce defensive repetition.
    set_text(
        find_paragraph(doc, "4.1 Desempenho alto"),
        "4.1 Recuperabilidade interna não é validade de construto",
    )
    set_text(
        find_paragraph(doc, "O resultado mais importante é simples") ,
        "O cross-fitting mostra que a regra geométrica reconstruída é recuperável fora da amostra, "
        "mesmo com ajuste confinado ao treino. A ROC-AUC média de 0,896 deve, portanto, ser lida como "
        "recuperabilidade interna do alvo endógeno. A avaliação responde corretamente à pergunta "
        "computacional definida pelo pipeline, mas essa pergunta não é equivalente a medir criminalidade.",
    )
    set_text(
        find_paragraph(doc, "4.2 O experimento como demonstração"),
        "4.2 A falácia da separabilidade endógena",
    )
    set_text(
        find_paragraph(doc, "A reconstrução atual funciona como") ,
        "Clustering, escolha do maior cluster, centroide, calibração, score e rótulo de teste "
        "permanecem relacionados à mesma geometria. Evitar vazamento direto melhora a validade interna, "
        "mas não cria um vínculo entre Y e um construto externo C. A falácia da separabilidade endógena "
        "surge quando a correção da operação g(Z) → Y é usada para sustentar uma alegação sobre C. O "
        "controle sintético torna essa distinção observável sem recorrer a propriedades humanas.",
    )
    set_text(
        find_paragraph(doc, "A análise de estabilidade examina") ,
        "Separabilidade e estabilidade respondem a perguntas diferentes. A primeira avalia quão bem a "
        "regra construída é recuperada; a segunda, quanto a própria regra depende de decisões "
        "computacionais. A validade de construto exige ainda evidência independente de que o alvo mede "
        "o conceito nomeado, e a validade externa exige população, seleção e condições de generalização "
        "definidas. O framework impede que sucesso em um plano seja usado como atalho para os demais.",
    )
    stability_discussion = find_paragraph(doc, "A auditoria final amplia substancialmente")
    set_text(
        stability_discussion,
        "As distribuições completas tornam a instabilidade um resultado, e não apenas um item de "
        "protocolo. ARI e Jaccard baixos aparecem em seeds, ordem, batch_size e representação, enquanto "
        "a prevalência varia de modo diferente. A escolha de k também altera a prevalência do maior "
        "cluster e a concordância das partições. Logo, seed, k, ordem, minibatch, representação e regra "
        "de seleção do alvo devem ser tratados como componentes da evidência, não detalhes neutros de "
        "implementação (von Luxburg, 2010).",
    )
    old_stability_discussion = find_paragraph(doc, "Os oito contrastes numericamente disponíveis")
    set_text(
        old_stability_discussion,
        "Os resultados permanecem condicionados à grade executada e ao MiniBatchKMeans. Eles não "
        "identificam uma partição verdadeira, mas rejeitam a apresentação de uma única solução como "
        "estrutura inevitável dos dados. Uma auditoria completa deve reportar distribuições, tamanho e "
        "prevalência do alvo, além das métricas de separabilidade.",
    )
    contribution_disc = find_paragraph(doc, "A contribuição de maior alcance")
    set_text(
        contribution_disc,
        "A contribuição de maior alcance é um método reutilizável para auditar construtos endógenos. "
        "A matriz de identificabilidade delimita as alegações possíveis; a linhagem localiza a origem "
        "do alvo; o cross-fitting testa recuperação sem ajuste no teste; controles negativos expõem a "
        "circularidade; e a estabilidade quantifica a dependência de escolhas computacionais. O caso de "
        "listas de procurados mostra por que a matemática do pipeline pode ser correta enquanto o nome "
        "atribuído à saída permanece cientificamente inválido.",
    )
    fairface_disc = find_paragraph(doc, "O braço FairFace amplia essa contribuição")
    set_text(
        fairface_disc,
        "O braço FairFace acrescenta uma análise de sensibilidade de composição, não uma validação "
        "racial do corpus principal. A coexistência de métricas agregadas semelhantes, partições "
        "diferentes e forte variação entre seeds demonstra contingência do pipeline. Replicações "
        "pareadas de amostragem são necessárias antes de atribuir variação adicional à composição.",
    )

    # Strengthen limitations and future-work boundaries.
    fairface_limit = find_paragraph(doc, "Braço demográfico externo.")
    set_text(
        fairface_limit,
        "Braço demográfico externo. Os rótulos do FairFace são categorias históricas do dataset, não "
        "autoidentificação irrestrita nem categorias biológicas. Como cada cenário altera proporções e "
        "registros, e a estabilidade entre seeds é baixa, não se estima efeito causal demográfico. Uma "
        "extensão adequada exige 50–100 replicações independentes ou um desenho pareado que maximize um "
        "núcleo comum, permitindo decompor composição, amostragem, k, seed e interações.",
    )
    internal_limit = find_paragraph(doc, "Validade interna e de conclusão.")
    set_text(
        internal_limit,
        "Validade interna e de conclusão. Group_id bloqueia duplicidades prováveis no limiar 0,999, mas "
        "não foi validado manualmente e não garante identidade única. A análise futura deve revisar "
        "cegamente pares próximos a 0,995, 0,997, 0,999 e 0,9995 e repetir métricas por limiar. Também "
        "deve ablar a regra do alvo — maior cluster, cluster aleatório, mais compacto e mais separado — "
        "para quantificar a influência dessa convenção. Os limiares de relevância atuais permanecem "
        "heurísticos; distribuições contínuas sustentam a inferência principal.",
    )
    reproduction_limit = find_paragraph(doc, "Reprodutibilidade e agenda seguinte.")
    set_text(
        reproduction_limit,
        "Reprodutibilidade e agenda seguinte. A corrida registra 78 testes aprovados, hashes de 20 "
        "outputs, configuração e sementes, mas worktree_dirty = true impede tratá-la como artefato "
        "canônico imutável. A próxima liberação deve partir de tag limpa, ambiente congelado, pesos "
        "identificados por SHA-256, manifesto completo e depósito dos outputs estruturados com DOI. "
        "Predições por observação devem ser preservadas para intervalos bootstrap auditáveis. Os "
        "experimentos adicionais são agenda futura e não são apresentados como resultados desta edição.",
    )

    # Concise conclusion centered on the framework rather than repeated disclaimers.
    conclusion_1 = find_paragraph(doc, "RQ1: o corpus observado")
    set_text(
        conclusion_1,
        "A auditoria demonstra três fatos complementares. Primeiro, o corpus é condicionado por "
        "seleção institucional e técnica e, portanto, não representa criminalidade populacional. "
        "Segundo, o alvo endógeno é internamente recuperável: em 9.482 embeddings, o cross-fitting "
        "produziu ROC-AUC média 0,896. Terceiro, essa recuperabilidade coexiste com baixa estabilidade: "
        "em k = 64, as medianas entre sementes foram 0,085 para ARI e 0,004 para Jaccard do alvo. O "
        "controle sintético reproduziu o mecanismo com AUC 1,000 e mudança completa do alvo após "
        "perturbação do clustering.",
    )
    conclusion_2 = find_paragraph(doc, "RQ4: a composição")
    set_text(
        conclusion_2,
        "O braço FairFace reforça a contingência sem sustentar causalidade racial: mudanças de "
        "composição e amostragem alteraram partições e conjuntos-alvo, enquanto prevalência e ROC-AUC "
        "variaram pouco, sobre forte instabilidade entre seeds. A contribuição geral é a distinção "
        "operacional entre separabilidade, estabilidade, validade de construto e validade externa. "
        "Quando Y é produzido a partir de Z e g(Z) recupera Y, o desempenho pode ser matematicamente "
        "real e substantivamente circular. Auditar essa passagem — localizando a origem do rótulo, "
        "quantificando sua estabilidade e exigindo evidência independente para o construto externo — "
        "é aplicável a qualquer sistema que converta representações de pessoas em alegações sobre "
        "caráter, risco, capacidade ou identidade social.",
    )

    # ABNT post-textual order and appendix terminology.
    appendix = find_paragraph(doc, "MATERIAL SUPLEMENTAR S1", "Heading 1")
    set_text(appendix, "APÊNDICE A — RASTREABILIDADE E CONTRATO DE REPRODUÇÃO")
    appendix.paragraph_format.page_break_before = True
    appendix_intro = find_paragraph(doc, "Este suplemento distingue")
    set_text(
        appendix_intro,
        "Este apêndice distingue o que a auditoria final efetivamente preencheu do que permanece "
        "histórico ou não recuperado. A classificação evita transformar documentação posterior em "
        "procedimento histórico e registra os elementos necessários para replicação independente.",
    )
    set_text(find_paragraph(doc, "Quadro S1"), "Quadro A.1 — Status do pacote de reprodução após a corrida final")
    set_text(find_paragraph(doc, "Quadro S2"), "Quadro A.2 — Especificação da configuração computacional da reconstrução final")
    set_text(find_paragraph(doc, "Quadro S3"), "Quadro A.3 — Escopo final da estabilidade e controle metodológico")
    references = find_paragraph(doc, "REFERÊNCIAS", "Heading 1")
    body = doc._body._element
    children = list(body)
    ref_index = children.index(references._p)
    ref_block = [node for node in children[ref_index:] if node.tag != qn("w:sectPr")]
    for node in ref_block:
        appendix._p.addprevious(node)
    references.paragraph_format.page_break_before = True

    # Add primary references in alphabetical order.
    bowyer = find_paragraph(doc, "BOWYER", "Reference")
    add_reference_before(
        bowyer,
        "BIRHANE, A.; DEHDASTIAN, S.; PRABHU, V. U.; BODDETI, V. The Dark Side of Dataset "
        "Scaling: Evaluating Racial Classification in Multimodal Models. In: ACM CONFERENCE ON "
        "FAIRNESS, ACCOUNTABILITY, AND TRANSPARENCY, 2024, Rio de Janeiro. Proceedings [...]. New "
        "York: ACM, 2024. DOI: 10.1145/3630106.3658968. Disponível em: "
        "https://doi.org/10.1145/3630106.3658968. Acesso em: 16 ago. 2026.",
    )
    kaufman = find_paragraph(doc, "KAUFMAN", "Reference")
    add_reference_before(
        kaufman,
        "KÄRKKÄINEN, K.; JOO, J. FairFace: Face Attribute Dataset for Balanced Race, Gender, "
        "and Age for Bias Measurement and Mitigation. In: IEEE/CVF WINTER CONFERENCE ON "
        "APPLICATIONS OF COMPUTER VISION, 2021. Proceedings [...]. p. 1548–1558. DOI: "
        "10.1109/WACV48630.2021.00159. Disponível em: https://doi.org/10.1109/WACV48630.2021.00159. "
        "Acesso em: 16 ago. 2026.",
    )
    pedregosa = find_paragraph(doc, "PEDREGOSA", "Reference")
    add_reference_before(
        pedregosa,
        "NIST. Face Recognition Technology Evaluation: Demographic Effects. Gaithersburg: "
        "National Institute of Standards and Technology, 2026. Disponível em: "
        "https://pages.nist.gov/frvt/html/frvt_demographics.html. Acesso em: 16 ago. 2026.",
    )
    roberts = find_paragraph(doc, "ROBERTS", "Reference")
    add_reference_before(
        roberts,
        "RAJI, I. D.; KUMAR, I. E.; HOROWITZ, A.; SELBST, A. D. The Fallacy of AI Functionality. "
        "In: ACM CONFERENCE ON FAIRNESS, ACCOUNTABILITY, AND TRANSPARENCY, 2022. Proceedings [...]. "
        "p. 959–972. DOI: 10.1145/3531146.3533158. Disponível em: "
        "https://doi.org/10.1145/3531146.3533158. Acesso em: 16 ago. 2026.",
    )

    # Update remaining title references and precision in body text.
    for paragraph in doc.paragraphs:
        text = paragraph.text
        updated = text.replace(
            "A falsa ideia de “criminalidade facial”: viés racial, listas públicas de procurados e os limites da similaridade facial: Auditoria sociotécnica de um pipeline que produz separabilidade sem medir criminalidade",
            TITLE_PT,
        )
        updated = updated.replace("0,896170", "0,896").replace("0,348207", "0,348").replace("0,043299", "0,043")
        updated = updated.replace("demonstração por construção", "demonstração construtiva")
        if updated != text:
            set_text(paragraph, updated)

    # Apply compact, legible geometry to all multi-column tables, especially the appendix.
    geometry_by_header = {
        "Alegação": ([1.3, 1.75, 2.35, 1.25], 7.8),
        "Transição": ([1.25, 2.15, 2.15, 1.1], 7.8),
        "Nível": ([1.35, 2.8, 2.2], 8.2),
        "Estágio/contagem": ([1.7, 0.8, 3.85], 8.1),
        "k": ([0.8, 1.25, 1.05, 1.35, 1.05, 1.3], 7.1),
        "Métrica": ([1.25, 2.75, 2.35], 8.0),
        "Cenário": ([0.85, 1.25, 1.0, 1.3, 1.0, 1.0], 7.6),
        "Componente": ([1.5, 3.0, 1.85], 7.5),
        "Elemento": ([1.4, 3.1, 1.85], 7.4),
        "Bloco": ([1.35, 3.1, 1.9], 7.8),
    }
    geometry_by_columns = {2: [1.0, 4.8], 3: [1.4, 3.0, 1.9], 4: [1.3, 1.8, 2.2, 1.1], 6: [0.6, 1.2, 1.0, 1.2, 1.0, 1.0]}
    for table in doc.tables:
        cols = len(table.columns)
        first_header = table.rows[0].cells[0].text.strip() if cols > 1 else ""
        if first_header in geometry_by_header:
            weights, size = geometry_by_header[first_header]
        else:
            weights = geometry_by_columns.get(cols, [1.0] * cols)
            size = 8.2 if len(table.rows) > 10 or cols >= 6 else 8.8
        format_table(table, weights, size)

    alt_texts = [
        "Diagrama em oito etapas: seleção institucional, publicação, coleta, detecção e alinhamento, embedding, clustering, escolha do alvo e avaliação; destaca os limites entre separabilidade interna e validade externa.",
        "Gráfico em quatro painéis contrastando ROC-AUC e PR-AUC por dobra com distribuições de ARI, Jaccard do alvo e prevalência do maior cluster ao longo de seis valores de k.",
        "Distribuição média das sete categorias fornecidas pelo FairFace dentro do cluster-alvo para os quatro cenários, em k igual a 64.",
        "Distribuições de ARI e Jaccard do alvo entre sementes para os quatro cenários do experimento FairFace, em k igual a 64.",
    ]
    for shape, alt in zip(doc.inline_shapes, alt_texts):
        doc_pr = shape._inline.docPr
        doc_pr.set("descr", alt)
        doc_pr.set("title", alt[:120])

    # Page behavior and academic paragraph hygiene.
    for paragraph in doc.paragraphs:
        if paragraph.style.name in {"Heading 1", "FrontHeading"}:
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name in {"Heading 2", "Heading 3", "TableCaption"}:
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name in {"Normal", "Keywords", "Reference", "Declaration"}:
            set_no_hyphenation(paragraph)
    # Empty styled remnants from the moved abstract block would otherwise create a blank page.
    for paragraph in list(doc.paragraphs):
        if not paragraph.text.strip() and paragraph.style.name in {"Heading 1", "Keywords"}:
            if not paragraph._p.xpath("./w:pPr/w:sectPr"):
                delete_paragraph(paragraph)
    summary_text.paragraph_format.keep_together = True
    abstract_text.paragraph_format.keep_together = True

    # Ask Word to refresh fields and the new TOC on open as an additional safeguard.
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    core = doc.core_properties
    core.title = TITLE_PT
    core.subject = "Auditoria sociotécnica de construtos endógenos em aprendizado de máquina"
    core.keywords = "validade de construto; viés de seleção; reconhecimento facial; auditoria algorítmica; fisiognomia algorítmica"
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Resumo words: {word_count(SUMMARY_PT)}")
    print(f"Abstract words: {word_count(ABSTRACT_EN)}")
    print(f"Paragraphs: {len(doc.paragraphs)}; tables: {len(doc.tables)}; figures: {len(doc.inline_shapes)}")


if __name__ == "__main__":
    build()
