from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


path = Path(
    r"C:\Users\Henrique Lima\Desktop\Repositórios\ia-uspJailer\docs\Listas públicas de procurados não medem criminalidade — revisão metodológica.docx"
)
document = Document(path)
for table in document.tables:
    properties = table._tbl.tblPr
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
document.save(path)

text = "\n".join(paragraph.text for paragraph in document.paragraphs)
text += "\n" + "\n".join(
    cell.text for table in document.tables for row in table.rows for cell in row.cells
)
print(
    {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "figures": len(document.inline_shapes),
        "toc_marker": text.count("[[TOC]]"),
        "old_title_fragment": text.count("A falsa ideia de"),
    }
)
